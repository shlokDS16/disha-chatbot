"""OCR wrapper around the OCR.space API.

OCR.space accepts PDFs and common image types (JPG, PNG, TIFF, BMP). For
.docx we extract text natively via python-docx — OCR.space does not take DOCX.

Design note: this is synchronous inside a thread-pool-friendly call; the
/files endpoints run it in a BackgroundTask so the upload response returns
immediately with status=queued.
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

import httpx
from docx import Document

from app.config import Settings
from app.errors import DishaError, ErrorCode
from app.models.schemas import Language

_log = logging.getLogger("disha.ocr")

_LANG_TO_OCR: dict[Language, str] = {
    Language.EN: "eng",
    Language.HI: "hin",
    Language.MR: "mar",
}


def _extract_docx(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


def _call_ocr_space(
    settings: Settings, path: Path, language: Language, is_pdf: bool
) -> str:
    if not settings.ocr_space_api_key:
        raise DishaError(
            ErrorCode.OCR_UNAVAILABLE,
            "OCR service is not configured on the server.",
            status_code=503,
        )
    data: dict[str, Any] = {
        "apikey": settings.ocr_space_api_key,
        "language": _LANG_TO_OCR.get(language, "eng"),
        "isOverlayRequired": "false",
        "OCREngine": "2",
        "scale": "true",
        "detectOrientation": "true",
    }
    if is_pdf:
        data["filetype"] = "PDF"

    with path.open("rb") as fh:
        files = {"file": (path.name, fh)}
        try:
            resp = httpx.post(
                settings.ocr_space_endpoint,
                data=data,
                files=files,
                timeout=120.0,
            )
        except httpx.HTTPError as e:
            _log.exception("ocr.space request failed")
            raise DishaError(
                ErrorCode.OCR_UNAVAILABLE,
                "OCR service failed to respond.",
                status_code=503,
                details={"reason": str(e)},
            ) from e

    if resp.status_code != 200:
        raise DishaError(
            ErrorCode.OCR_UNAVAILABLE,
            f"OCR service returned HTTP {resp.status_code}.",
            status_code=503,
        )
    body = resp.json()
    if body.get("IsErroredOnProcessing"):
        err_msg = body.get("ErrorMessage") or ["Unknown OCR error"]
        raise DishaError(
            ErrorCode.OCR_UNAVAILABLE,
            "OCR failed to parse the document.",
            status_code=422,
            details={"ocr_error": err_msg},
        )
    parsed = body.get("ParsedResults") or []
    return "\n".join((r.get("ParsedText") or "").strip() for r in parsed).strip()


def run_ocr(settings: Settings, path: str, ftype: str, language: Language) -> str:
    """Extract text from the file. ftype is one of pdf/docx/image."""
    p = Path(path)
    if not p.exists():
        raise DishaError(
            ErrorCode.FILE_NOT_FOUND,
            "Uploaded file missing from storage.",
            status_code=404,
            details={"path": str(p)},
        )
    if ftype == "docx":
        return _extract_docx(p)
    return _call_ocr_space(settings, p, language, is_pdf=(ftype == "pdf"))
